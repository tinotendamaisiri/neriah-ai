# shared/document_extractor.py
# Text extraction from all document types submitted by tertiary students.
# Detects document type and routes to the correct extraction method.
# First step in the tertiary grading pipeline.

from __future__ import annotations

import io
import logging

from .ocr_client import analyse_image

logger = logging.getLogger(__name__)


def detect_document_type(file_bytes: bytes, filename: str) -> str:
    """Detect document type from filename extension and file magic bytes.

    Returns one of: "pdf", "pdf_scanned", "docx", "image", "mixed", "unknown"
    """
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("jpg", "jpeg", "png", "tiff", "tif", "bmp", "webp"):
        return "image"

    if ext in ("doc",):
        logger.warning(
            "Old .doc format is not supported for file '%s'. "
            "Please ask the student to convert to .docx.",
            filename,
        )
        return "unknown"

    if ext == "docx":
        return "docx"

    if ext == "pdf":
        return _classify_pdf(file_bytes)

    # No recognised extension — fall back to magic bytes
    magic = file_bytes[:8]
    if magic[:3] == b"\xff\xd8\xff":
        return "image"
    if magic[:4] == b"\x89PNG":
        return "image"
    if magic[:4] in (b"II*\x00", b"MM\x00*"):
        return "image"
    if magic[:2] == b"BM":
        return "image"
    if magic[:4] == b"RIFF":
        return "image"
    if magic[:4] == b"%PDF":
        return _classify_pdf(file_bytes)

    return "unknown"


def _classify_pdf(file_bytes: bytes) -> str:
    """Return 'pdf', 'pdf_scanned', or 'mixed' by inspecting each page."""
    import pdfplumber  # noqa: PLC0415 — lazy import to reduce cold-start time
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_with_text = 0
            pages_without_text = 0
            for page in pdf.pages:
                text = page.extract_text() or ""
                if len(text) > 20:
                    pages_with_text += 1
                else:
                    pages_without_text += 1

        if pages_with_text == 0:
            return "pdf_scanned"
        if pages_without_text == 0:
            return "pdf"
        return "mixed"
    except Exception:
        logger.exception("Failed to classify PDF; defaulting to 'pdf_scanned'")
        return "pdf_scanned"


async def extract_text(
    file_bytes: bytes,
    filename: str,
    document_type: str | None = None,
) -> tuple[str, str]:
    """Extract text from a document, returning (text, document_type).

    If document_type is None, it is detected automatically.
    """
    if document_type is None:
        document_type = detect_document_type(file_bytes, filename)

    if document_type == "pdf":
        text = await _extract_pdf_text(file_bytes)
    elif document_type == "pdf_scanned":
        text = await _extract_via_ocr(file_bytes)
    elif document_type == "docx":
        text = await _extract_docx_text(file_bytes)
    elif document_type == "image":
        text = await _extract_via_ocr(file_bytes)
    elif document_type == "mixed":
        text = await _extract_mixed_pdf(file_bytes)
    else:
        raise ValueError(f"Unsupported document type for file: {filename}")

    logger.info(
        "Extracted %d characters from '%s' (type=%s)",
        len(text),
        filename,
        document_type,
    )
    return text, document_type


async def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from a native PDF using pdfplumber."""
    import pdfplumber  # noqa: PLC0415
    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            pages.append(page_text.strip())
    return "\n\n".join(p for p in pages if p)


async def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a .docx Word document."""
    from docx import Document as DocxDocument  # noqa: PLC0415
    doc = DocxDocument(io.BytesIO(file_bytes))

    parts: list[str] = []

    # Paragraphs
    parts.append("\n".join(p.text for p in doc.paragraphs))

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return "\n".join(p for p in parts if p)


async def _extract_via_ocr(file_bytes: bytes) -> str:
    """Extract text from an image or scanned PDF page via Azure Document Intelligence."""
    logger.info("Using OCR extraction")
    text, _bounding_box = await analyse_image(file_bytes)
    return text


async def _extract_mixed_pdf(file_bytes: bytes) -> str:
    """Extract text from a mixed PDF (some native pages, some scanned)."""
    import pdfplumber  # noqa: PLC0415
    pages: list[str] = []
    text_page_count = 0
    ocr_page_count = 0

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if len(page_text) >= 20:
                pages.append(page_text.strip())
                text_page_count += 1
            else:
                # Convert page to JPEG and run OCR
                pil_image = page.to_image(resolution=150).original
                buf = io.BytesIO()
                pil_image.save(buf, format="JPEG")
                ocr_text = await _extract_via_ocr(buf.getvalue())
                pages.append(ocr_text)
                ocr_page_count += 1

    logger.info(
        "Mixed PDF extraction complete: %d text pages, %d OCR pages",
        text_page_count,
        ocr_page_count,
    )
    return "\n\n".join(p for p in pages if p)
